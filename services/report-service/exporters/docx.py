import os
from docx import Document

class DOCXExporter:
    @staticmethod
    def export(markdown_text: str, output_path: str):
        """
        Converts markdown text to a simple DOCX using python-docx.
        """
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        
        document = Document()
        
        for line in markdown_text.split('\n'):
            line = line.strip()
            if not line:
                continue
            elif line.startswith('# '):
                document.add_heading(line[2:], 0)
            elif line.startswith('## '):
                document.add_heading(line[3:], 1)
            elif line.startswith('- '):
                document.add_paragraph(line[2:], style='List Bullet')
            elif line.startswith('1. ') or line.startswith('2. '):
                document.add_paragraph(line[3:], style='List Number')
            else:
                document.add_paragraph(line)
                
        document.save(output_path)
        return output_path
